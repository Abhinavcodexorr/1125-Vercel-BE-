from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = 'http://localhost:3002/api/v1'
PROD_BASE = 'https://api1125.vercel.app/api/v1'


def set_cell_shading(cell, fill):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_api_heading(doc, num, name):
    p = doc.add_paragraph()
    run = p.add_run(f'{num}. {name}')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()


def add_field(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    p.add_run(value)


def add_request_table(doc, rows):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Request:-')
    r.bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['Parameter Name', 'Required/Optional', 'Data Type', 'Description']
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(hdr[i], 'D9E2F3')

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_response(doc, json_text):
    p = doc.add_paragraph()
    r = p.add_run('Response:-')
    r.bold = True
    code = doc.add_paragraph(json_text)
    code.style = 'No Spacing'
    for run in code.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_auth_note(doc, auth_type='Admin JWT'):
    add_field(doc, 'Auth', auth_type)
    add_field(doc, 'Header', 'x-access-token: <token>  OR  Authorization: Bearer <token>')


def add_public_note(doc):
    add_field(doc, 'Auth', 'None — public website API')


def build_doc():
    doc = Document()
    doc.add_paragraph()
    add_title(doc, 'API Documentation')
    doc.add_paragraph()
    add_field(doc, 'Project', '1125 Hotel / Palm Island Backend')
    add_field(doc, 'Base URL (Local)', BASE)
    add_field(doc, 'Base URL (Production)', PROD_BASE)
    add_field(doc, 'Version', 'v1')
    doc.add_paragraph()

    # Section header
    sec = doc.add_paragraph()
    sec_run = sec.add_run('ADMIN APIs')
    sec_run.bold = True
    sec_run.font.size = Pt(16)
    doc.add_paragraph()

    n = 1

    # 1. Admin Login
    add_api_heading(doc, n, 'Admin Login API')
    add_field(doc, 'Description', 'This API is used to login SuperAdmin, Manager, or SubAdmin and receive JWT token.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/login')
    add_field(doc, 'Method', 'POST')
    add_request_table(doc, [
        ['email', 'required', 'string', 'Admin email address'],
        ['password', 'required', 'string', 'Admin password'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Login successful",
  "data": {
    "token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",
    "user": {
      "_id": "...",
      "name": "Admin",
      "email": "1125demo@gmail.com",
      "role": "SuperAdmin"
    }
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 2. Get Current Admin
    add_api_heading(doc, n, 'Get Current Admin API')
    add_field(doc, 'Description', 'This API returns the logged-in admin user profile.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/me')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['(none)', 'optional', '-', 'No body parameters. Send JWT in header.'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "User retrieved successfully",
  "data": {
    "_id": "...",
    "name": "Admin",
    "email": "1125demo@gmail.com",
    "role": "SuperAdmin"
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 3. Update Password
    add_api_heading(doc, n, 'Update Admin Password API')
    add_field(doc, 'Description', 'This API is used to change the logged-in admin password.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/password')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['oldPassword', 'required', 'string', 'Current password'],
        ['newPassword', 'required', 'string', 'New password'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Password updated successfully",
  "data": null,
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 4. Admin Logout
    add_api_heading(doc, n, 'Admin Logout API')
    add_field(doc, 'Description', 'This API logs out the current SuperAdmin session.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/logout')
    add_field(doc, 'Method', 'POST')
    add_auth_note(doc, 'SuperAdmin JWT')
    add_request_table(doc, [
        ['(none)', 'optional', '-', 'No body parameters.'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Logout successful",
  "data": null,
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 5. Create SubAdmin
    add_api_heading(doc, n, 'Create SubAdmin API')
    add_field(doc, 'Description', 'This API creates a new SubAdmin user. Only SuperAdmin can create.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/subadmin')
    add_field(doc, 'Method', 'POST')
    add_auth_note(doc, 'SuperAdmin JWT')
    add_request_table(doc, [
        ['name', 'required', 'string', 'SubAdmin full name'],
        ['email', 'required', 'string', 'SubAdmin email'],
        ['password', 'required', 'string', 'SubAdmin password'],
        ['role', 'optional', 'string', 'Default: SubAdmin'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 201,
  "message": "SubAdmin created successfully",
  "data": {
    "_id": "...",
    "name": "John Manager",
    "email": "manager@hotel.com",
    "role": "SubAdmin"
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 6. List SubAdmins
    add_api_heading(doc, n, 'List SubAdmins API')
    add_field(doc, 'Description', 'This API returns all SubAdmin / Manager users.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/subadmin')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['(none)', 'optional', '-', 'No query parameters.'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "SubAdmins retrieved successfully",
  "data": [
    {
      "_id": "...",
      "name": "John Manager",
      "email": "manager@hotel.com",
      "role": "SubAdmin",
      "isBlocked": false
    }
  ],
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 7. Get SubAdmin By ID
    add_api_heading(doc, n, 'Get SubAdmin By ID API')
    add_field(doc, 'Description', 'This API returns a single SubAdmin by ID.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/subadmin/:id')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'SubAdmin MongoDB _id'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "SubAdmin retrieved successfully",
  "data": {
    "_id": "...",
    "name": "John Manager",
    "email": "manager@hotel.com",
    "role": "SubAdmin"
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 8. Update SubAdmin
    add_api_heading(doc, n, 'Update SubAdmin API')
    add_field(doc, 'Description', 'This API updates SubAdmin details. SubAdmin cannot update others.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/subadmin/:id')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'SubAdmin MongoDB _id'],
        ['name', 'optional', 'string', 'Updated name'],
        ['email', 'optional', 'string', 'Updated email'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "SubAdmin updated successfully",
  "data": { "_id": "...", "name": "Updated Name", "email": "updated@hotel.com" },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 9. Block SubAdmin
    add_api_heading(doc, n, 'Block SubAdmin API')
    add_field(doc, 'Description', 'This API blocks a SubAdmin from logging in.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/subadmin/:id/block')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'SubAdmin MongoDB _id'],
        ['isBlocked', 'optional', 'boolean', 'Set true to block'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "SubAdmin blocked successfully",
  "data": { "_id": "...", "isBlocked": true },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 10. Unblock SubAdmin
    add_api_heading(doc, n, 'Unblock SubAdmin API')
    add_field(doc, 'Description', 'This API unblocks a SubAdmin so they can login again.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/subadmin/:id/unblock')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'SubAdmin MongoDB _id'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "SubAdmin unblocked successfully",
  "data": { "_id": "...", "isBlocked": false },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 11. Delete SubAdmin
    add_api_heading(doc, n, 'Delete SubAdmin API')
    add_field(doc, 'Description', 'This API permanently deletes a SubAdmin user.')
    add_field(doc, 'Endpoint', '/api/v1/superadmin/subadmin/:id')
    add_field(doc, 'Method', 'DELETE')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'SubAdmin MongoDB _id'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "SubAdmin deleted successfully",
  "data": null,
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # --- ROOMS ADMIN ---
    sec2 = doc.add_paragraph()
    sec2_run = sec2.add_run('ADMIN — ROOMS APIs')
    sec2_run.bold = True
    sec2_run.font.size = Pt(14)
    doc.add_paragraph()

    # 12. List Rooms Admin
    add_api_heading(doc, n, 'List All Rooms (Admin) API')
    add_field(doc, 'Description', 'This API returns all rooms including inactive. Excludes soft-deleted rooms.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/admin')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['isActive', 'optional', 'string', 'Filter: true or false (query param)'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Rooms retrieved successfully",
  "data": {
    "total": 2,
    "data": [
      {
        "_id": "...",
        "title": "The Villa",
        "slug": "the-villa",
        "type": "Villa",
        "price": 500,
        "currency": "GHS",
        "guests": 5,
        "size": 850,
        "unit": "sq ft",
        "amenities": [],
        "images": [{ "url": "https://...", "order": 0 }],
        "blockedDates": [],
        "isActive": true,
        "isDeleted": false
      }
    ]
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 13. Get Room Admin
    add_api_heading(doc, n, 'Get Room By ID (Admin) API')
    add_field(doc, 'Description', 'This API returns a single room by MongoDB _id or slug.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/admin/:idOrSlug')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug e.g. the-villa'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room retrieved successfully",
  "data": {
    "_id": "...",
    "title": "The Villa",
    "slug": "the-villa",
    "type": "Villa",
    "price": 500,
    "size": 850,
    "blockedDates": []
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 14. Create Room
    add_api_heading(doc, n, 'Create Room API')
    add_field(doc, 'Description', 'This API creates a new room/villa in the hotel catalog.')
    add_field(doc, 'Endpoint', '/api/v1/rooms')
    add_field(doc, 'Method', 'POST')
    add_auth_note(doc)
    add_request_table(doc, [
        ['title', 'required', 'string', 'Room name e.g. Presidential Villa'],
        ['type', 'required', 'string', 'Room type e.g. Villa, Suite'],
        ['price', 'required', 'number', 'Price per night'],
        ['guests', 'required', 'number', 'Max guests (min 1)'],
        ['quantity', 'optional', 'number', 'Number of identical units. Default: 1'],
        ['size', 'required', 'number', 'Room size value'],
        ['slug', 'optional', 'string', 'URL slug (auto-generated from title if omitted)'],
        ['description', 'optional', 'string', 'Room description'],
        ['currency', 'optional', 'string', 'Default: GHS'],
        ['unit', 'optional', 'string', 'Size unit. Default: sq ft'],
        ['amenities', 'optional', 'array', '[{ name, icon, iconType, key }]'],
        ['images', 'optional', 'array', '[{ url, order }] — use upload API first'],
        ['isActive', 'optional', 'boolean', 'Default: true'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 201,
  "message": "Room created successfully",
  "data": {
    "_id": "...",
    "title": "Presidential Villa",
    "slug": "presidential-villa",
    "type": "Villa",
    "price": 500,
    "guests": 4,
    "size": 850,
    "unit": "sq ft",
    "images": [],
    "isActive": true
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 15. Update Room
    add_api_heading(doc, n, 'Update Room API')
    add_field(doc, 'Description', 'This API updates an existing room. Send only fields to change.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
        ['title', 'optional', 'string', 'Updated title'],
        ['type', 'optional', 'string', 'Updated type'],
        ['price', 'optional', 'number', 'Updated price per night'],
        ['guests', 'optional', 'number', 'Updated max guests'],
        ['quantity', 'optional', 'number', 'Updated number of units (cannot be less than booked units)'],
        ['size', 'optional', 'number', 'Updated size'],
        ['description', 'optional', 'string', 'Updated description'],
        ['amenities', 'optional', 'array', 'Updated amenities list'],
        ['images', 'optional', 'array', 'Updated images list'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room updated successfully",
  "data": { "_id": "...", "title": "Updated Villa", "price": 550 },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 16. Room Status
    add_api_heading(doc, n, 'Activate / Deactivate Room API')
    add_field(doc, 'Description', 'This API activates or deactivates a room on the website.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug/status')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
        ['isActive', 'required', 'boolean', 'true = show on website, false = hide'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room activated successfully",
  "data": { "_id": "...", "isActive": true },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 17. Delete Room
    add_api_heading(doc, n, 'Delete Room API')
    add_field(doc, 'Description', 'This API soft-deletes a room. Room is hidden and marked isDeleted true.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug')
    add_field(doc, 'Method', 'DELETE')
    add_auth_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room deleted successfully",
  "data": { "_id": "...", "slug": "the-villa" },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 18. Upload Room Image
    add_api_heading(doc, n, 'Upload Room Image API')
    add_field(doc, 'Description', 'This API uploads one or more room images to AWS S3 bucket (rooms/ folder).')
    add_field(doc, 'Endpoint', '/api/v1/upload/room-image')
    add_field(doc, 'Method', 'POST')
    add_auth_note(doc)
    add_request_table(doc, [
        ['files', 'required', 'file (multipart)', 'Image file(s). Field name must be "files"'],
        ['order', 'optional', 'number', 'Starting display order. Default: 0'],
    ])
    add_field(doc, 'Content-Type', 'multipart/form-data')
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room image(s) uploaded successfully",
  "data": {
    "total": 1,
    "images": [
      {
        "url": "https://palmisland.s3.eu-north-1.amazonaws.com/rooms/photo_123.png",
        "file_name": "photo_123.png",
        "order": 0
      }
    ]
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 19. Block Room Dates
    add_api_heading(doc, n, 'Block Room Dates API')
    add_field(doc, 'Description', 'This API blocks dates for a room so guests cannot book (maintenance, private events).')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug/blocked-dates')
    add_field(doc, 'Method', 'POST')
    add_auth_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
        ['startDate', 'required', 'string', 'Block start date YYYY-MM-DD (first blocked night)'],
        ['endDate', 'required', 'string', 'Block end date YYYY-MM-DD (checkout day, not blocked)'],
        ['reason', 'optional', 'string', 'Reason e.g. Maintenance'],
        ['blocks', 'optional', 'array', 'Multiple blocks: [{ startDate, endDate, reason }]'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room dates blocked successfully",
  "data": {
    "room": { "_id": "...", "title": "The Villa", "slug": "the-villa" },
    "added": [
      {
        "startDate": "2026-07-01T00:00:00.000Z",
        "endDate": "2026-07-05T00:00:00.000Z",
        "reason": "Maintenance",
        "nights": 4
      }
    ],
    "blockedDates": ["2026-07-01", "2026-07-02", "2026-07-03", "2026-07-04"]
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 20. Get Blocked Dates
    add_api_heading(doc, n, 'Get Room Blocked Dates API')
    add_field(doc, 'Description', 'This API returns all admin-blocked date ranges for a room.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug/blocked-dates')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Blocked dates retrieved successfully",
  "data": {
    "room": { "_id": "...", "title": "The Villa", "slug": "the-villa" },
    "total": 1,
    "blocked": [
      {
        "_id": "...",
        "startDate": "2026-07-01T00:00:00.000Z",
        "endDate": "2026-07-05T00:00:00.000Z",
        "reason": "Maintenance",
        "occupiedDates": ["2026-07-01", "2026-07-02", "2026-07-03", "2026-07-04"]
      }
    ],
    "blockedDates": ["2026-07-01", "2026-07-02", "2026-07-03", "2026-07-04"]
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 21. Unblock Room Dates
    add_api_heading(doc, n, 'Unblock Room Dates API')
    add_field(doc, 'Description', 'This API removes one blocked date range from a room.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug/blocked-dates/:blockId')
    add_field(doc, 'Method', 'DELETE')
    add_auth_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
        ['blockId (path)', 'required', 'string', 'Blocked range _id from blockedDates array'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room dates unblocked successfully",
  "data": {
    "removed": {
      "_id": "...",
      "startDate": "2026-07-01T00:00:00.000Z",
      "endDate": "2026-07-05T00:00:00.000Z",
      "reason": "Maintenance"
    },
    "blockedDates": []
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # --- BOOKING ADMIN ---
    sec3 = doc.add_paragraph()
    sec3_run = sec3.add_run('ADMIN — BOOKING APIs')
    sec3_run.bold = True
    sec3_run.font.size = Pt(14)
    doc.add_paragraph()

    # 22. List Bookings
    add_api_heading(doc, n, 'List All Bookings (Admin) API')
    add_field(doc, 'Description', 'This API returns all bookings for admin panel.')
    add_field(doc, 'Endpoint', '/api/v1/booking')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['filter', 'optional', 'string', 'Query: incomplete | paid | cancelled'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Bookings retrieved successfully",
  "data": [
    {
      "_id": "...",
      "bookingReference": "ABC12345",
      "checkInDate": "2026-06-05T00:00:00.000Z",
      "checkOutDate": "2026-06-08T00:00:00.000Z",
      "status": "Confirmed",
      "paymentStatus": "paid",
      "totalAmount": 1500
    }
  ],
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 23. Update Booking Status
    add_api_heading(doc, n, 'Update Booking Status API')
    add_field(doc, 'Description', 'This API updates booking status (Checked-In, Checked-Out, Cancelled, etc.).')
    add_field(doc, 'Endpoint', '/api/v1/booking/:id/status')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'Booking MongoDB _id'],
        ['status', 'required', 'string', 'Pending | Confirmed | Checked-In | Checked-Out | Cancelled | No-Show'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Booking status updated successfully",
  "data": { "_id": "...", "status": "Checked-In" },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 24. Manual Confirm Booking
    add_api_heading(doc, n, 'Manual Confirm Booking API')
    add_field(doc, 'Description', 'This API manually confirms a booking and marks payment as paid.')
    add_field(doc, 'Endpoint', '/api/v1/booking/:id/manual-confirm')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'Booking MongoDB _id'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Booking confirmed successfully",
  "data": { "_id": "...", "status": "Confirmed", "paymentStatus": "paid" },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 25. Cancel Booking
    add_api_heading(doc, n, 'Cancel Booking API')
    add_field(doc, 'Description', 'This API cancels a booking and sends cancellation email.')
    add_field(doc, 'Endpoint', '/api/v1/booking/:id/cancel')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'Booking MongoDB _id'],
        ['cancellationReason', 'optional', 'string', 'Reason for cancellation'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Booking cancelled successfully",
  "data": {
    "_id": "...",
    "status": "Cancelled",
    "cancellationReason": "Guest request",
    "cancelledAt": "2026-05-26T10:00:00.000Z"
  },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # --- PROMO ADMIN ---
    sec4 = doc.add_paragraph()
    sec4_run = sec4.add_run('ADMIN — PROMO CODE APIs')
    sec4_run.bold = True
    sec4_run.font.size = Pt(14)
    doc.add_paragraph()

    # 26. Create Promo
    add_api_heading(doc, n, 'Create Promo Code API')
    add_field(doc, 'Description', 'This API creates a new discount promo code.')
    add_field(doc, 'Endpoint', '/api/v1/promo')
    add_field(doc, 'Method', 'POST')
    add_auth_note(doc)
    add_request_table(doc, [
        ['code', 'required', 'string', 'Promo code e.g. SUMMER10'],
        ['discountType', 'required', 'string', 'percentage or flat'],
        ['discountValue', 'required', 'number', 'Discount amount or percentage'],
        ['startDate', 'optional', 'date', 'Promo valid from'],
        ['endDate', 'optional', 'date', 'Promo valid until'],
        ['isActive', 'optional', 'boolean', 'Default: true'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 201,
  "message": "Promo code created successfully",
  "data": { "_id": "...", "code": "SUMMER10", "discountValue": 10, "isActive": true },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 27. List Promo
    add_api_heading(doc, n, 'List Promo Codes API')
    add_field(doc, 'Description', 'This API returns all promo codes for admin.')
    add_field(doc, 'Endpoint', '/api/v1/promo/admin')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc)
    add_request_table(doc, [
        ['(none)', 'optional', '-', 'No parameters'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Promo codes retrieved successfully",
  "data": [{ "_id": "...", "code": "SUMMER10", "isActive": true }],
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 28. Update Promo
    add_api_heading(doc, n, 'Update Promo Code API')
    add_field(doc, 'Description', 'This API updates an existing promo code.')
    add_field(doc, 'Endpoint', '/api/v1/promo/:id')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'Promo MongoDB _id'],
        ['code', 'optional', 'string', 'Updated code'],
        ['discountValue', 'optional', 'number', 'Updated discount'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Promo code updated successfully",
  "data": { "_id": "...", "code": "SUMMER15" },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 29. Promo Status
    add_api_heading(doc, n, 'Activate / Deactivate Promo API')
    add_field(doc, 'Description', 'This API activates or deactivates a promo code.')
    add_field(doc, 'Endpoint', '/api/v1/promo/:id/status')
    add_field(doc, 'Method', 'PUT')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'Promo MongoDB _id'],
        ['isActive', 'required', 'boolean', 'true or false'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Promo code activated successfully",
  "data": { "_id": "...", "isActive": true },
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # 30. Delete Promo
    add_api_heading(doc, n, 'Delete Promo Code API')
    add_field(doc, 'Description', 'This API deletes a promo code.')
    add_field(doc, 'Endpoint', '/api/v1/promo/:id')
    add_field(doc, 'Method', 'DELETE')
    add_auth_note(doc)
    add_request_table(doc, [
        ['id (path)', 'required', 'string', 'Promo MongoDB _id'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Promo code deleted successfully",
  "data": null,
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # --- CONTACT ADMIN ---
    sec5 = doc.add_paragraph()
    sec5_run = sec5.add_run('ADMIN — CONTACT APIs')
    sec5_run.bold = True
    sec5_run.font.size = Pt(14)
    doc.add_paragraph()

    # 31. List Contact Messages
    add_api_heading(doc, n, 'List Contact Messages API')
    add_field(doc, 'Description', 'This API returns all contact form messages submitted from website.')
    add_field(doc, 'Endpoint', '/api/v1/contact')
    add_field(doc, 'Method', 'GET')
    add_auth_note(doc, 'SuperAdmin JWT')
    add_request_table(doc, [
        ['(none)', 'optional', '-', 'No parameters'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Messages retrieved successfully",
  "data": [
    {
      "_id": "...",
      "name": "John Doe",
      "email": "john@email.com",
      "message": "I want to book a villa",
      "createdAt": "2026-05-26T10:00:00.000Z"
    }
  ],
  "timestamp": "2026-05-26T10:00:00.000Z"
}''')
    n += 1

    # --- WEBSITE / PUBLIC ---
    sec_web = doc.add_paragraph()
    sec_web_run = sec_web.add_run('WEBSITE — PUBLIC APIs')
    sec_web_run.bold = True
    sec_web_run.font.size = Pt(16)
    doc.add_paragraph()
    add_field(doc, 'Note', 'All website APIs are public (no JWT). Same room routes also work at /api/v1/cabin')

    # W1. List Rooms
    add_api_heading(doc, n, 'List Rooms (Website) API')
    add_field(doc, 'Description', 'Returns active rooms for the booking website. Filters by stay dates, adults, and quantity. Only rooms available for the requested stay are returned when dates are provided.')
    add_field(doc, 'Endpoint', '/api/v1/rooms')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['page', 'optional', 'number', 'Page number. Default: 1'],
        ['limit', 'optional', 'number', 'Items per page (max 50). Default: 10'],
        ['checkInDate', 'optional', 'string', 'Stay check-in YYYY-MM-DD'],
        ['checkOutDate', 'optional', 'string', 'Stay check-out YYYY-MM-DD'],
        ['adult / adults', 'optional', 'number', 'Number of adults — filters by room capacity'],
        ['children / child', 'optional', 'number', 'Number of children. Default: 0'],
        ['quantity / qty / units', 'optional', 'number', 'Units needed — only rooms with enough free units are returned'],
    ])
    add_field(doc, 'Example', f'{BASE}/rooms?page=1&limit=10&checkInDate=2026-06-08&checkOutDate=2026-06-09&adult=2&quantity=1')
    add_response(doc, '''{
  "success": true,
  "totalItems": 1,
  "page": 1,
  "limit": 10,
  "data": [
    {
      "_id": "...",
      "name": "The Villa",
      "title": "5-Bedroom Beach Residence",
      "slug": "the-villa",
      "description": "...",
      "size": 450,
      "unit": "sq m",
      "pricePerNight": 650,
      "currency": "USD",
      "currencySymbol": "$",
      "formattedPrice": "$ 650.00/night",
      "guests": 10,
      "quantity": 2,
      "adultCapacity": 10,
      "childCapacity": 0,
      "amenities": [],
      "images": [{ "url": "https://...", "alt": "The Villa", "order": 0 }],
      "availability": {
        "isAvailable": true,
        "quantity": 2,
        "availableUnits": 2,
        "bookedUnits": 0,
        "requestedQuantity": 1,
        "nights": 1,
        "subTotal": 650,
        "unavailableReason": null,
        "blockedDates": []
      }
    }
  ]
}''')
    n += 1

    # W2. Get Room By ID
    add_api_heading(doc, n, 'Get Room By ID (Website) API')
    add_field(doc, 'Description', 'Returns single room details with availability for optional stay dates.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room MongoDB _id or slug e.g. the-villa'],
        ['checkInDate', 'optional', 'string', 'Check-in YYYY-MM-DD (query)'],
        ['checkOutDate', 'optional', 'string', 'Check-out YYYY-MM-DD (query)'],
        ['adult / adults', 'optional', 'number', 'Adults count (query)'],
        ['quantity', 'optional', 'number', 'Units needed (query)'],
    ])
    add_response(doc, '''{
  "success": true,
  "data": {
    "_id": "...",
    "name": "The Villa",
    "slug": "the-villa",
    "pricePerNight": 650,
    "quantity": 2,
    "availability": {
      "isAvailable": true,
      "availableUnits": 2,
      "requestedQuantity": 1,
      "nights": 1,
      "subTotal": 650
    }
  }
}''')
    n += 1

    # W3. Check availability before checkout
    add_api_heading(doc, n, 'Check Room Stay Availability API')
    add_field(doc, 'Description', 'Pre-payment availability check for checkout page. Validates dates, adults, and quantity before booking.')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug/check-availability')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
        ['checkInDate', 'required', 'string', 'Check-in YYYY-MM-DD'],
        ['checkOutDate', 'required', 'string', 'Check-out YYYY-MM-DD'],
        ['adult / adults', 'required', 'number', 'Number of adults'],
        ['children', 'optional', 'number', 'Number of children'],
        ['quantity', 'optional', 'number', 'Units to book. Default: 1'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 200,
  "message": "Room is available for selected dates",
  "data": {
    "roomId": "...",
    "slug": "the-villa",
    "name": "The Villa",
    "checkInDate": "2026-06-08",
    "checkOutDate": "2026-06-09",
    "adults": 2,
    "children": 0,
    "nights": 1,
    "quantity": 2,
    "requestedQuantity": 1,
    "availableUnits": 2,
    "bookedUnits": 0,
    "pricePerNight": 650,
    "subTotal": 650,
    "totalAmount": 650,
    "currency": "USD",
    "isAvailable": true
  }
}''')
    n += 1

    # W4. Room availability calendar
    add_api_heading(doc, n, 'Get Room Availability Calendar API')
    add_field(doc, 'Description', 'Returns all booked, blocked, and available dates for a room (no date range params).')
    add_field(doc, 'Endpoint', '/api/v1/rooms/:idOrSlug/availability')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['idOrSlug (path)', 'required', 'string', 'Room _id or slug'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Room availability retrieved successfully",
  "data": {
    "room": { "_id": "...", "title": "The Villa", "quantity": 2 },
    "bookedDates": ["2026-06-10"],
    "blockedDates": ["2026-07-01"],
    "availableDates": ["2026-06-08", "2026-06-09"],
    "partiallyBookedDates": [],
    "occupancyByDate": {
      "2026-06-10": { "bookedCount": 1, "availableUnits": 1, "quantity": 2 }
    }
  }
}''')
    n += 1

    # --- CART ---
    sec_cart = doc.add_paragraph()
    sec_cart_run = sec_cart.add_run('WEBSITE — CART APIs')
    sec_cart_run.bold = True
    sec_cart_run.font.size = Pt(14)
    doc.add_paragraph()

    # W5. Add to cart
    add_api_heading(doc, n, 'Add Room To Cart API')
    add_field(doc, 'Description', 'Adds a room to cart after checking availability for the requested quantity and dates. Returns cartId — save it on the frontend.')
    add_field(doc, 'Endpoint', '/api/v1/cart/add')
    add_field(doc, 'Method', 'POST')
    add_public_note(doc)
    add_request_table(doc, [
        ['roomId', 'required', 'string', 'Room MongoDB _id'],
        ['checkInDate', 'required', 'string', 'Check-in YYYY-MM-DD'],
        ['checkOutDate', 'required', 'string', 'Check-out YYYY-MM-DD'],
        ['adults / adult', 'required', 'number', 'Number of adults (min 1)'],
        ['children / child', 'optional', 'number', 'Number of children. Default: 0'],
        ['quantity', 'required', 'number', 'Units to book (min 1). Checked against room availability'],
        ['cartId', 'optional', 'string', 'Existing cart ID to append item. Omit to create new cart'],
    ])
    add_field(doc, 'Content-Type', 'application/json')
    add_response(doc, '''{
  "success": true,
  "statusCode": 201,
  "message": "Room added to cart",
  "data": {
    "cartId": "454bb0cf-4bf3-4e0f-90f0-dfdf03d66cab",
    "subTotal": 650,
    "currency": "USD",
    "expiresAt": "2026-06-12T12:00:00.000Z",
    "allAvailable": true,
    "items": [
      {
        "_id": "...",
        "roomId": "...",
        "roomSnapshot": {
          "title": "The Villa",
          "slug": "the-villa",
          "price": 650,
          "currency": "USD",
          "guests": 10,
          "quantity": 2
        },
        "checkInDate": "2026-06-08T00:00:00.000Z",
        "checkOutDate": "2026-06-09T00:00:00.000Z",
        "adults": 2,
        "children": 0,
        "quantity": 1,
        "nights": 1,
        "pricePerNight": 650,
        "subTotal": 650,
        "isAvailable": true,
        "unavailableReason": null
      }
    ]
  }
}''')
    n += 1

    # W6. Get cart
    add_api_heading(doc, n, 'Get Cart API')
    add_field(doc, 'Description', 'Returns cart with all items. Re-checks availability on each request.')
    add_field(doc, 'Endpoint', '/api/v1/cart/:cartId')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['cartId (path)', 'required', 'string', 'Cart ID from add-to-cart response'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Cart retrieved successfully",
  "data": {
    "cartId": "454bb0cf-4bf3-4e0f-90f0-dfdf03d66cab",
    "subTotal": 650,
    "currency": "USD",
    "allAvailable": true,
    "items": []
  }
}''')
    n += 1

    # W7. Check cart availability
    add_api_heading(doc, n, 'Check Cart Availability API')
    add_field(doc, 'Description', 'Call on checkout page before payment. Re-validates all cart items including quantity.')
    add_field(doc, 'Endpoint', '/api/v1/cart/:cartId/check-availability')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['cartId (path)', 'required', 'string', 'Cart ID'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Cart availability checked",
  "data": {
    "cartId": "...",
    "allAvailable": true,
    "subTotal": 650,
    "items": [{ "isAvailable": true, "quantity": 1, "unavailableReason": null }]
  }
}''')
    n += 1

    # W8. Remove cart item
    add_api_heading(doc, n, 'Remove Cart Item API')
    add_field(doc, 'Description', 'Removes one item from the cart.')
    add_field(doc, 'Endpoint', '/api/v1/cart/:cartId/items/:itemId')
    add_field(doc, 'Method', 'DELETE')
    add_public_note(doc)
    add_request_table(doc, [
        ['cartId (path)', 'required', 'string', 'Cart ID'],
        ['itemId (path)', 'required', 'string', 'Cart item _id from items array'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Cart item removed",
  "data": { "cartId": "...", "subTotal": 0, "items": [] }
}''')
    n += 1

    # W9. Clear cart
    add_api_heading(doc, n, 'Clear Cart API')
    add_field(doc, 'Description', 'Removes all items from the cart.')
    add_field(doc, 'Endpoint', '/api/v1/cart/:cartId')
    add_field(doc, 'Method', 'DELETE')
    add_public_note(doc)
    add_request_table(doc, [
        ['cartId (path)', 'required', 'string', 'Cart ID'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Cart cleared",
  "data": { "cartId": "...", "subTotal": 0, "items": [] }
}''')
    n += 1

    # --- BOOKING WEBSITE ---
    sec_book = doc.add_paragraph()
    sec_book_run = sec_book.add_run('WEBSITE — BOOKING & PAYMENT (Hubtel)')
    sec_book_run.bold = True
    sec_book_run.font.size = Pt(14)
    doc.add_paragraph()

    # W10. Create booking
    add_api_heading(doc, n, 'Create Room Booking API (Hubtel)')
    add_field(doc, 'Description', 'Creates booking from cart or direct room selection. Checks quantity availability, then initiates Hubtel payment. Redirect user to checkoutUrl.')
    add_field(doc, 'Endpoint', '/api/v1/booking/room')
    add_field(doc, 'Method', 'POST')
    add_public_note(doc)
    add_request_table(doc, [
        ['cartId', 'optional*', 'string', 'Cart ID — use this OR roomId below'],
        ['roomId', 'optional*', 'string', 'Direct booking without cart'],
        ['checkInDate', 'required**', 'string', 'Required when using roomId (not cart)'],
        ['checkOutDate', 'required**', 'string', 'Required when using roomId'],
        ['adults', 'required**', 'number', 'Required when using roomId'],
        ['quantity', 'required**', 'number', 'Units to book when using roomId'],
        ['children', 'optional', 'number', 'Children count'],
        ['guestDetails', 'required', 'object', 'Guest information object'],
        ['guestDetails.firstName', 'required', 'string', 'Guest first name'],
        ['guestDetails.lastName', 'required', 'string', 'Guest last name'],
        ['guestDetails.email', 'required', 'string', 'Guest email'],
        ['guestDetails.mobileNumber', 'required', 'string', 'Guest phone e.g. 233241234567'],
        ['guestDetails.countryCode', 'optional', 'string', 'Country code'],
        ['guestDetails.specialRequests', 'optional', 'string', 'Special requests'],
    ])
    add_field(doc, 'Note', '* Send cartId OR roomId. ** Required only for direct roomId booking.')
    add_field(doc, 'Content-Type', 'application/json')
    add_response(doc, '''{
  "success": true,
  "statusCode": 201,
  "message": "Booking created. Complete payment with Hubtel.",
  "data": {
    "bookingReference": "ABC12XYZ",
    "bookingIds": ["..."],
    "totalAmount": 650,
    "currency": "USD",
    "paymentMethod": "Hubtel",
    "checkoutUrl": "https://pay.hubtel.com/...",
    "bookings": []
  }
}''')
    n += 1

    # W11. Hubtel callback
    add_api_heading(doc, n, 'Hubtel Payment Callback API')
    add_field(doc, 'Description', 'Server-to-server webhook from Hubtel when payment completes. Configure HUBTEL_CALLBACK_URL to this endpoint.')
    add_field(doc, 'Endpoint', '/api/v1/booking/hubtel/callback')
    add_field(doc, 'Method', 'POST')
    add_field(doc, 'Auth', 'Called by Hubtel servers')
    add_request_table(doc, [
        ['ClientReference', 'required', 'string', 'bookingReference sent during checkout'],
        ['Status', 'required', 'string', 'Paid | Unpaid | etc.'],
        ['TransactionId', 'optional', 'string', 'Hubtel transaction ID'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Callback processed"
}''')
    n += 1

    # W12. Confirm payment
    add_api_heading(doc, n, 'Confirm Booking Payment API')
    add_field(doc, 'Description', 'Called after customer returns from Hubtel checkout page to verify payment status.')
    add_field(doc, 'Endpoint', '/api/v1/booking/confirm')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['reference', 'required', 'string', 'bookingReference query param'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Booking payment status checked",
  "data": {
    "isPaid": true,
    "status": "Paid",
    "bookings": [{ "bookingReference": "ABC12XYZ", "status": "Confirmed", "paymentStatus": "paid" }]
  }
}''')
    n += 1

    # W13. Get booking by reference
    add_api_heading(doc, n, 'Get Booking By Reference API')
    add_field(doc, 'Description', 'Returns booking details by booking reference number.')
    add_field(doc, 'Endpoint', '/api/v1/booking/reference/:reference')
    add_field(doc, 'Method', 'GET')
    add_public_note(doc)
    add_request_table(doc, [
        ['reference (path)', 'required', 'string', 'Booking reference e.g. ABC12XYZ'],
    ])
    add_response(doc, '''{
  "success": true,
  "message": "Booking retrieved",
  "data": {
    "bookingReference": "ABC12XYZ",
    "bookings": [{
      "roomId": "...",
      "roomSnapshot": { "title": "The Villa" },
      "roomQuantity": 1,
      "checkInDate": "2026-06-08T00:00:00.000Z",
      "checkOutDate": "2026-06-09T00:00:00.000Z",
      "totalAmount": 650,
      "status": "Confirmed",
      "paymentStatus": "paid"
    }]
  }
}''')
    n += 1

    # W14. Contact submit
    add_api_heading(doc, n, 'Submit Contact Message API')
    add_field(doc, 'Description', 'Public contact form submission from website.')
    add_field(doc, 'Endpoint', '/api/v1/contact')
    add_field(doc, 'Method', 'POST')
    add_public_note(doc)
    add_request_table(doc, [
        ['name', 'required', 'string', 'Sender name'],
        ['email', 'required', 'string', 'Sender email'],
        ['message', 'required', 'string', 'Message or query text'],
    ])
    add_response(doc, '''{
  "success": true,
  "statusCode": 201,
  "message": "Message sent successfully",
  "data": { "_id": "...", "name": "John", "email": "john@email.com" }
}''')

    out_admin = r'd:\1125 BE\1125 BE\docs\1125_API_Documentation_Admin.docx'
    out_full = r'd:\1125 BE\1125 BE\docs\1125_API_Documentation.docx'
    doc.save(out_admin)
    doc.save(out_full)
    print(f'Saved: {out_admin}')
    print(f'Saved: {out_full}')


if __name__ == '__main__':
    build_doc()
