"""
Generate a Word document with email template screenshots and editable content
for team review. Run from repo root: python scripts/generate-email-templates-docx.py
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = ROOT / "templates"
DOCS_DIR = ROOT / "docs"
SCREENSHOTS_DIR = DOCS_DIR / "email-template-screenshots"
OUTPUT_DOCX = DOCS_DIR / "1125-Beach-Villa-Email-Templates-Review.docx"

BRAND_RGB = RGBColor(0x5A, 0x8A, 0xAD)
SAMPLE_DATE = datetime.now().strftime("%B %d, %Y at %I:%M %p")

EMAIL_TEMPLATES = [
    {
        "id": "contact-enquiry",
        "title": "1. Contact Enquiry (Enquiry Raised)",
        "html": TEMPLATES_DIR / "contact-enquiry" / "index.html",
        "subject": "Enquiry Raised",
        "to": "info@1125beachvilla.com",
        "bcc": None,
        "from_name": "1125 Beach Villa",
        "from_email": "noreply@1125beachvilla.com",
        "trigger": "Sent when a visitor submits the Contact Us / enquiry form.",
        "audience": "Internal — villa team only (not sent to the guest).",
        "header": "Enquiry Raised",
        "sections": [
            {
                "heading": "Contact Details",
                "fields": [
                    ("Name", "{{guest_name}}"),
                    ("Email", "{{guest_email}}"),
                    ("Date", "{{submitted_at}}"),
                ],
            },
            {
                "heading": "Message",
                "body": "{{enquiry_message}}",
            },
        ],
        "footer": [
            "1125 Beach Villa",
            "This is an automated email from the contact enquiry form.",
            "© {{year}} 1125 Beach Villa. All rights reserved.",
        ],
        "sample": {
            "guest_name": "John Doe",
            "guest_email": "john.doe@example.com",
            "submitted_at": SAMPLE_DATE,
            "enquiry_message": "Hello, I would like to enquire about availability at 1125 Beach Villa for next month.",
            "year": str(datetime.now().year),
        },
    },
    {
        "id": "booking-in-progress",
        "title": "2. Booking In Progress",
        "html": TEMPLATES_DIR / "booking-in-progress" / "index.html",
        "subject": "Booking In Progress - Reference: {{booking_reference}}",
        "to": "info@1125beachvilla.com",
        "bcc": None,
        "from_name": "1125 Beach Villa",
        "from_email": "noreply@1125beachvilla.com",
        "trigger": "Sent when a new booking is created (payment not yet completed).",
        "audience": "Internal — villa team only (not sent to the guest).",
        "header": "Booking In Progress",
        "sections": [
            {
                "heading": "Guest Details",
                "fields": [
                    ("Name", "{{guest_name}}"),
                    ("Email", "{{guest_email}}"),
                    ("Phone", "{{guest_phone}}"),
                    ("Date", "{{submitted_at}}"),
                ],
            },
            {
                "heading": "Booking Details",
                "fields": [
                    ("Reference", "{{booking_reference}}"),
                    ("Check-in", "{{check_in}}"),
                    ("Check-out", "{{check_out}}"),
                    ("Room", "{{room_name}}"),
                    ("Guests", "{{guests}}"),
                    ("Total Amount", "{{total_amount}}"),
                ],
            },
        ],
        "footer": [
            "1125 Beach Villa",
            "This is an automated email from the booking system.",
            "© {{year}} 1125 Beach Villa. All rights reserved.",
        ],
        "sample": {
            "guest_name": "John Doe",
            "guest_email": "john.doe@example.com",
            "guest_phone": "+233 24 000 0000",
            "submitted_at": SAMPLE_DATE,
            "booking_reference": "BK-2026-00482",
            "check_in": "Fri, Jun 27, 2026 · after 2:00 PM",
            "check_out": "Sun, Jun 29, 2026 · before 11:00 AM",
            "room_name": "Ocean View Suite",
            "guests": "2 adults",
            "total_amount": "GHS 4,500.00",
            "year": str(datetime.now().year),
        },
    },
    {
        "id": "booking-confirmation",
        "title": "3. Booking Confirmation",
        "html": TEMPLATES_DIR / "booking-confirmation" / "index.html",
        "subject": "Booking Confirmation - Reference: {{booking_reference}}",
        "to": "{{guest_email}}",
        "bcc": None,
        "from_name": "1125 Beach Villa",
        "from_email": "noreply@1125beachvilla.com",
        "trigger": "Sent to the guest when payment is successful and the booking is confirmed.",
        "audience": "Guest (customer).",
        "header": "Booking Confirmed",
        "intro": "Hi {{guest_name}},\n\nThank you for booking with 1125 Beach Villa. Your reservation is confirmed — we look forward to welcoming you.",
        "sections": [
            {
                "heading": "Booking Details",
                "fields": [
                    ("Reference", "{{booking_reference}}"),
                    ("Check-in", "{{check_in}}"),
                    ("Check-out", "{{check_out}}"),
                    ("Room", "{{room_name}}"),
                    ("Guests", "{{guests}}"),
                    ("Amount Paid", "{{amount_paid}}"),
                ],
            },
        ],
        "notes": [
            "If you have any questions about your stay, please contact us at info@1125beachvilla.com.",
        ],
        "footer": [
            "1125 Beach Villa",
            "This is an automated booking confirmation email.",
            "© {{year}} 1125 Beach Villa. All rights reserved.",
        ],
        "sample": {
            "guest_name": "John Doe",
            "guest_email": "john.doe@example.com",
            "booking_reference": "BK-2026-00482",
            "check_in": "Fri, Jun 27, 2026 · after 2:00 PM",
            "check_out": "Sun, Jun 29, 2026 · before 11:00 AM",
            "room_name": "Ocean View Suite",
            "guests": "2 adults",
            "amount_paid": "GHS 4,500.00",
            "year": str(datetime.now().year),
        },
    },
    {
        "id": "booking-cancellation",
        "title": "4. Booking Cancellation",
        "html": TEMPLATES_DIR / "booking-cancellation" / "index.html",
        "subject": "Booking Cancellation - Reference: {{booking_reference}}",
        "to": "{{guest_email}}",
        "bcc": "info@1125beachvilla.com",
        "from_name": "1125 Beach Villa",
        "from_email": "noreply@1125beachvilla.com",
        "trigger": "Sent when an admin cancels a booking.",
        "audience": "Guest (customer), with BCC copy to info@1125beachvilla.com.",
        "header": "Booking Cancelled",
        "intro": "Hi {{guest_name}},\n\nYour booking with 1125 Beach Villa has been cancelled. Details are below.",
        "sections": [
            {
                "heading": "Cancelled Booking",
                "fields": [
                    ("Reference", "{{booking_reference}}"),
                    ("Room", "{{room_name}}"),
                    ("Check-in", "{{check_in}}"),
                    ("Check-out", "{{check_out}}"),
                    ("Guests", "{{guests}}"),
                    ("Total Amount", "{{total_amount}}"),
                    ("Refund Amount", "{{refund_amount}}"),
                ],
            },
        ],
        "notes": [
            "{{refund_policy_message}}",
            "If you have any questions about this cancellation, please contact us at info@1125beachvilla.com.",
        ],
        "footer": [
            "1125 Beach Villa",
            "This is an automated booking cancellation email.",
            "© {{year}} 1125 Beach Villa. All rights reserved.",
        ],
        "sample": {
            "guest_name": "John Doe",
            "guest_email": "john.doe@example.com",
            "booking_reference": "BK-2026-00482",
            "room_name": "Ocean View Suite",
            "check_in": "Fri, Jun 27, 2026",
            "check_out": "Sun, Jun 29, 2026",
            "guests": "2 adults",
            "total_amount": "GHS 4,500.00",
            "refund_amount": "GHS 4,500.00",
            "refund_policy_message": "As your booking was cancelled 7 days or more before the check-in date, you are eligible for a full refund. Refunds are processed within 14 business days.",
            "year": str(datetime.now().year),
        },
    },
]


def fill(template: str, values: dict) -> str:
    result = template
    for key, value in values.items():
        result = result.replace(f"{{{{{key}}}}}", str(value))
    return result


def capture_screenshots() -> dict[str, Path]:
    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={"width": 720, "height": 900})

        for tpl in EMAIL_TEMPLATES:
            html_path = tpl["html"]
            if not html_path.exists():
                raise FileNotFoundError(f"Missing template HTML: {html_path}")

            file_url = html_path.resolve().as_uri()
            page.goto(file_url, wait_until="networkidle")
            page.wait_for_timeout(500)

            target = page.locator(".email-container")
            screenshot_path = SCREENSHOTS_DIR / f"{tpl['id']}.png"
            target.screenshot(path=str(screenshot_path))
            paths[tpl["id"]] = screenshot_path
            print(f"Screenshot: {screenshot_path}")

        browser.close()

    return paths


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = BRAND_RGB


def add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)


def render_sample_content(tpl: dict) -> str:
    sample = tpl["sample"]
    lines = [f"HEADER: {tpl['header']}", ""]

    if tpl.get("intro"):
        lines.append(fill(tpl["intro"], sample))
        lines.append("")

    for section in tpl.get("sections", []):
        lines.append(section["heading"].upper())
        if section.get("fields"):
            for label, value_tpl in section["fields"]:
                lines.append(f"{label}: {fill(value_tpl, sample)}")
        if section.get("body"):
            lines.append(fill(section["body"], sample))
        lines.append("")

    for note in tpl.get("notes", []):
        lines.append(fill(note, sample))
        lines.append("")

    lines.append("FOOTER")
    for line in tpl.get("footer", []):
        lines.append(fill(line, sample))

    return "\n".join(lines).strip()


def build_document(screenshots: dict[str, Path]) -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("1125 Beach Villa — Email Templates Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BRAND_RGB

    intro = doc.add_paragraph(
        "This document contains visual previews and editable copy for all transactional "
        "email templates. Update the text in the “Email content (editable)” sections below "
        "and share your feedback. Placeholders in {{double_braces}} are replaced with real "
        "data when emails are sent."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run("Generated: ").bold = True
    meta.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
    doc.add_paragraph()

    doc.add_heading("Brand & delivery defaults", level=1)
    add_label_value(doc, "From name", "1125 Beach Villa")
    add_label_value(doc, "From email", "noreply@1125beachvilla.com")
    add_label_value(doc, "Internal inbox", "info@1125beachvilla.com")
    add_label_value(doc, "Primary colour", "#5a8aad")
    doc.add_page_break()

    for tpl in EMAIL_TEMPLATES:
        add_heading(doc, tpl["title"], level=1)

        doc.add_heading("Delivery", level=2)
        add_label_value(doc, "Subject", fill(tpl["subject"], tpl["sample"]))
        add_label_value(doc, "To", fill(tpl["to"], tpl["sample"]))
        if tpl.get("bcc"):
            add_label_value(doc, "BCC", tpl["bcc"])
        add_label_value(doc, "From", f"{tpl['from_name']} <{tpl['from_email']}>")
        add_label_value(doc, "When sent", tpl["trigger"])
        add_label_value(doc, "Audience", tpl["audience"])
        add_label_value(doc, "Source file", str(tpl["html"].relative_to(ROOT)))

        doc.add_heading("Visual preview", level=2)
        shot = screenshots.get(tpl["id"])
        if shot and shot.exists():
            doc.add_picture(str(shot), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph("[Screenshot unavailable]")

        doc.add_heading("Email content (editable)", level=2)
        content_para = doc.add_paragraph(render_sample_content(tpl))
        for run in content_para.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)

        doc.add_heading("Placeholders used", level=2)
        blob = render_sample_content(tpl) + " " + tpl["subject"] + " " + tpl["to"]
        placeholders = sorted(set(re.findall(r"\{\{(\w+)\}\}", blob)))
        doc.add_paragraph(", ".join(f"{{{{{p}}}}}" for p in placeholders) if placeholders else "None")

        doc.add_heading("Review notes / requested changes", level=2)
        doc.add_paragraph(
            "Use this space to note wording, layout, or field changes:\n\n"
            "• Subject line:\n"
            "• Header title:\n"
            "• Body copy:\n"
            "• Footer:\n"
            "• Other:"
        )

        doc.add_page_break()

    doc.save(str(OUTPUT_DOCX))
    print(f"Document saved: {OUTPUT_DOCX}")


def main() -> None:
    screenshots = capture_screenshots()
    build_document(screenshots)


if __name__ == "__main__":
    main()
